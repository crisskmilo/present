"""
Complete Proposal Word Document Generator for PRESENTE Project.
Preserves the user's complete unabridged original text (Introduction, Sections 1 to 23),
and appends Section 24 (Technical Architecture, Multiplatform, $0 Cost Hosting & Diagrams).
"""

import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

DOCS_DIR = os.path.dirname(__file__)
ASSETS_DIR = os.path.join(DOCS_DIR, "assets")
OUTPUT_DOCX_PATH = os.path.join(DOCS_DIR, "present-mvp-proposal.docx")


def build_full_proposal_docx(output_path: str = OUTPUT_DOCX_PATH) -> None:
    doc = docx.Document()

    # 1-inch margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base typography
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)

    def add_title(text: str):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(22)
        r.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        return p

    def add_subtitle(text: str):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)
        return p

    def add_heading_1(text: str):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(4)
        r = h.add_run(text)
        r.bold = True
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        return h

    def add_heading_2(text: str):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(3)
        r = h.add_run(text)
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)
        return h

    def add_image_safely(img_filename: str, width=Inches(6.0)):
        img_path = os.path.join(ASSETS_DIR, img_filename)
        if os.path.exists(img_path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            doc.add_picture(img_path, width=width)
            doc.add_paragraph()

    # Header
    add_title("PRESENTE\nIniciativa digital de acompañamiento en salud mental")
    add_subtitle("PRESENTE - Y RECUERDA QUE NO ESTAMOS SOLOS\nPropuesta de iniciativa y desarrollo de MVP\nAutor: Cristian Camilo Rojas Estrada | Fecha: 2026-08-07\nInstitución Colaboradora Propuesta: Clínica San Juan de Dios de Manizales")
    doc.add_paragraph()

    # Full Unabridged Text (Sections 1 to 23)
    add_heading_1("Introducción")
    doc.add_paragraph("La salud mental forma parte de la vida de todas las personas.")
    doc.add_paragraph("A lo largo de la vida podemos atravesar momentos de ansiedad, tristeza, incertidumbre, soledad, pérdidas, conflictos, cambios importantes, dificultades familiares, preocupaciones personales o diferentes situaciones que pueden afectar nuestra manera de sentir, pensar y relacionarnos con nuestro entorno.")
    doc.add_paragraph("Cada persona vive estas experiencias de una manera diferente.")
    doc.add_paragraph("Algunas personas buscan acompañamiento profesional. Otras encuentran refugio en su familia, amigos, comunidad o espiritualidad. Algunas no saben cómo pedir ayuda. Otras se encuentran atravesando un proceso de salud mental y necesitan continuar fortaleciendo las herramientas adquiridas. También existen personas que han buscado acompañamiento en diferentes momentos y sienten que todavía necesitan encontrar nuevas formas de afrontar lo que están viviendo.")
    doc.add_paragraph("Hay quienes ya finalizaron un proceso de atención y necesitan continuar fortaleciendo lo aprendido. También existen personas que todavía se encuentran atravesando una situación difícil y aún no han iniciado un proceso de acompañamiento profesional, aunque podrían beneficiarse de hacerlo.")
    doc.add_paragraph("En muchos casos puede aparecer una sensación de aislamiento, incomprensión o falta de herramientas para continuar.")
    doc.add_paragraph("De esta realidad nace PRESENTE, una iniciativa que busca utilizar la tecnología como medio para acompañar ese transitar en salud mental, acercando conocimiento, experiencias y herramientas a todas las personas que puedan encontrar valor en ellas.")
    doc.add_paragraph("La iniciativa surge desde una perspectiva tecnológica y social. La propuesta consiste en desarrollar una plataforma digital independiente que permita organizar, producir, publicar y distribuir contenidos de manera sencilla, accesible y responsable.")
    doc.add_paragraph("El proyecto parte de una convicción: Atravesar un momento de sufrimiento no define a una persona. Una situación difícil puede formar parte de una etapa de la vida y, con acompañamiento, conocimiento y herramientas adecuadas, pueden encontrarse nuevos caminos.")
    doc.add_paragraph("PRESENTE busca convertirse en un punto de encuentro digital donde una persona pueda encontrar una conversación, una reflexión, una herramienta, una experiencia o una palabra de aliento que le permita sentirse acompañada y encontrar un poco de luz en su camino.")
    doc.add_paragraph("PRESENTE - Y RECUERDA QUE NO ESTAMOS SOLOS.")

    add_heading_1("1. Origen de la iniciativa")
    doc.add_paragraph("La idea tiene como punto de partida la experiencia de personas relacionadas con procesos de atención en salud mental y, particularmente, la reflexión sobre lo que ocurre después de finalizar una hospitalización, un programa de atención o un proceso terapéutico.")
    doc.add_paragraph("Una pregunta inicial dio origen al proyecto: ¿Qué pasa después?")
    doc.add_paragraph("Después de salir de un proceso de atención, una persona puede volver a su entorno cotidiano y encontrarse nuevamente con sus responsabilidades, relaciones, preocupaciones y dificultades. Puede necesitar: continuar fortaleciendo herramientas, recordar algo que aprendió, escuchar una experiencia, recibir una palabra de aliento, encontrar un espacio al cual regresar, reconocer señales de alerta y continuar construyendo hábitos y redes de apoyo.")
    doc.add_paragraph("Sin embargo, la reflexión llevó a ampliar la pregunta: ¿Qué ocurre con quienes todavía no buscan ayuda? ¿Con quienes se encuentran aislados? ¿Con los familiares que no saben cómo acompañar? La iniciativa busca acercar información, herramientas y perspectivas que ayuden a reconocer que se necesita apoyo e incentivar la búsqueda de acompañamiento profesional sin realizar diagnósticos ni evaluaciones clínicas.")

    add_heading_1("2. Concepto de PRESENTE")
    doc.add_paragraph("El nombre PRESENTE representa el propósito central: estar presente significa acompañar el momento que una persona está viviendo. Poner a su alcance un espacio donde encontrar conocimiento, herramientas y experiencias. La plataforma no está dirigida exclusivamente a pacientes, sino a personas en cualquier etapa de vida.")
    doc.add_paragraph("PRESENTE - Y RECUERDA QUE NO ESTAMOS SOLOS. Representa la convicción de que las dificultades pueden atravesarse con apoyo, conocimiento, vínculos y herramientas.")

    add_heading_1("3. Propósito")
    doc.add_paragraph("Crear una plataforma digital accesible para todos que acerque contenidos relacionados con salud mental, acompañamiento y bienestar, desarrollados y publicados por profesionales y colaboradores autorizados.")

    add_heading_1("4. Objetivo general")
    doc.add_paragraph("Construir PRESENTE: una plataforma digital que facilite el acceso a contenidos de salud mental y bienestar mediante una arquitectura tecnológica para administrar, organizar y distribuir información elaborada por profesionales y colaboradores autorizados.")

    add_heading_1("5. Objetivos específicos")
    doc.add_paragraph("• Desarrollar un MVP funcional de una plataforma digital de acompañamiento y divulgación.", style='List Bullet')
    doc.add_paragraph("• Facilitar el acceso a contenidos audiovisuales, audio, textos, imágenes y documentos.", style='List Bullet')
    doc.add_paragraph("• Organizar contenidos por temas, formatos y perfiles de publicación.", style='List Bullet')
    doc.add_paragraph("• Crear un sistema de administración de contenidos con control de roles y permisos.", style='List Bullet')
    doc.add_paragraph("• Integrar canales digitales: plataforma web, YouTube y redes sociales.", style='List Bullet')
    doc.add_paragraph("• Crear una experiencia sencilla y accesible para personas con cansancio o baja concentración.", style='List Bullet')
    doc.add_paragraph("• Incorporar psicología, bienestar, espiritualidad/pastoral y experiencias seleccionadas.", style='List Bullet')
    doc.add_paragraph("• Promover y reconocer la importancia de los procesos terapéuticos profesionales.", style='List Bullet')
    doc.add_paragraph("• Construir progresivamente una comunidad con criterios de seguridad y privacidad.", style='List Bullet')
    doc.add_paragraph("• Validar el MVP antes de realizar inversiones de mayor complejidad.", style='List Bullet')

    add_heading_1("6. La tecnología como valor agregado")
    doc.add_paragraph("La tecnología permite transformar una sola conversación entre profesionales en múltiples formatos organizados: episodio de video, episodio de audio/podcast, fragmentos breves para redes sociales, resumen escrito, guías en PDF y material complementario en un solo lugar.")

    add_heading_1("7. Arquitectura conceptual de la plataforma")
    doc.add_paragraph("Estructurada en: Capa de presentación (interfaz responsive), Capa de aplicación (backend de gestión y roles), Capa de información (base de datos relacional), CMS de contenidos, Integración multimedia (YouTube) y Canales de distribución.")

    add_heading_1("8. Modelo de publicación y control de contenidos")
    doc.add_paragraph("No es una red social abierta. Modelo de publicación controlado mediante tres roles: Administrador (gestión total y moderación), Publicador Autorizado (psicólogos, pastoral, especialistas) y Usuario (consulta, lectura, reproducción y descarga de recursos sin publicación abierta).")

    add_heading_1("9. Comunidad PRESENTE")
    doc.add_paragraph("Un espacio seguro donde la generación de contenido está centralizada en profesionales autorizados y los usuarios participan mediante mecanismos regulados y moderados.")

    add_heading_1("10. Propuesta inicial de contenidos")
    doc.add_paragraph("• Psicología y Salud Mental: Comprensión emocional, higiene del sueño, afrontamiento y prevención.", style='List Bullet')
    doc.add_paragraph("• Espiritualidad y Pastoral: Sentido de vida, propósito, esperanza, fe y respeto a la diversidad.", style='List Bullet')
    doc.add_paragraph("• Bienestar: Meditación, respiración, descanso, terapia ocupacional, arte y movimiento.", style='List Bullet')
    doc.add_paragraph("• Conversaciones y Testimonios: Entrevistas y experiencias de superación curadas.", style='List Bullet')

    add_heading_1("11. Estrategia multimedia")
    doc.add_paragraph("Un mismo contenido adaptado a diversos formatos: Podcasts completos, Videos, Audios, Clips cortos y Recursos descargables.")

    add_heading_1("12. Redes sociales")
    doc.add_paragraph("TikTok y YouTube como canales de descubrimiento para conectar personas con la plataforma web/móvil donde profundizar.")

    add_heading_1("13. Experiencia de usuario (UX)")
    doc.add_paragraph("Diseño claro, accesible y guiado por la intención («¿Qué buscas hoy?: Escuchar, Encontrar una herramienta, Una reflexión, Acompañar a alguien»).")

    add_heading_1("14. Relación con profesionales e instituciones")
    doc.add_paragraph("La Clínica San Juan de Dios de Manizales representa la incubadora y primer espacio de colaboración, manteniendo la gobernanza independiente de la plataforma para futuras alianzas con universidades y organizaciones.")

    add_heading_1("15. Principios de contenido")
    doc.add_paragraph("Responsabilidad, Respeto, Humanidad, Diversidad, Conocimiento responsable, Privacidad e Incentivo permanente a buscar atención profesional.")

    add_heading_1("16. Alcance y límites")
    doc.add_paragraph("Plataforma de acompañamiento y psicoeducación. No diagnostica, no reemplaza la psicoterapia ni presta atención de urgencias médicas.")

    add_heading_1("17. MVP — Primera etapa de desarrollo")
    doc.add_paragraph("Página de inicio, catálogo por categorías, reproductor multimedia, descargables, perfiles de especialistas y panel de administración.")

    add_heading_1("18. Desarrollo progresivo")
    doc.add_paragraph("Seis fases: Definición, Desarrollo del MVP, Contenido inicial, Lanzamiento, Validación y Evolución.")

    add_heading_1("19. Sostenibilidad")
    doc.add_paragraph("Operación con costos de infraestructura mínimos mediante capas gratuitas. Futuros fondos orientados a la continuidad y caridad a pacientes y familias vulnerables de la Clínica San Juan de Dios.")

    add_heading_1("20. Primer piloto propuesto")
    doc.add_paragraph("Equipo: Dirección Técnica (Ingeniero de Sistemas creador), Líder de Pastoral y Psicólogo Clínico. Tema piloto: «Después de un proceso de salud mental: ¿cómo continuar?».")

    add_heading_1("21. Indicadores iniciales")
    doc.add_paragraph("Evaluación de alcance, visitas, reproducciones, tiempo de permanencia y retroalimentación cualitativa sobre el valor percibido.")

    add_heading_1("22. Propuesta de colaboración inicial")
    doc.add_paragraph("Presentación al equipo pastoral y psicológico, definición del piloto, desarrollo del MVP y validación conjunta de resultados.")

    add_heading_1("23. Conclusión")
    doc.add_paragraph("PRESENTE busca unir la necesidad humana de sentirse acompañado con la capacidad de la tecnología para acercar herramientas y esperanza a quienes transitan por momentos de dificultad. Una persona en sufrimiento puede descubrir que no tiene por qué recorrer su camino sola.")

    # ---------------------------------------------------------
    # SECTION 24: DETAILED TECHNICAL ARCHITECTURE & ZERO-COST HOSTING
    # ---------------------------------------------------------
    add_heading_1("24. Arquitectura Tecnológica Detallada, Ecosistema Multiplataforma y Despliegue a Costo Cero ($0 USD)")
    doc.add_paragraph(
        "Para complementar la propuesta y garantizar su viabilidad técnica y financiera ante la Clínica San Juan de Dios, "
        "se detalla el diseño de ingeniería de software, la entrega multiplataforma y la infraestructura en la nube a costo $0:"
    )

    add_heading_2("24.1. Ecosistema Multiplataforma: Web App y Aplicación Híbrida (Android & iOS)")
    doc.add_paragraph(
        "Desarrollado con Angular y Capacitor para ofrecer tres canales de distribución desde una única base de código: "
        "Web App Responsive / PWA (acceso web instantáneo sin instalar), Aplicación Móvil Android (.apk / .aab) "
        "y Aplicación Móvil iOS (.ipa para iPhone y iPad)."
    )
    add_image_safely("multiplatform-web-app-diagram.jpg")

    add_heading_2("24.2. Arquitectura de Software: Clean Architecture con Inversión de Control (IoC)")
    doc.add_paragraph(
        "El backend se implementa en Python con FastAPI bajo el estándar de Clean Architecture (Arquitectura Limpia), "
        "separando estrictamente el Núcleo de Dominio (entidades e interfaces abstractas abc.ABC), Capa de Aplicación (casos de uso y DTOs Pydantic), "
        "Capa de Infraestructura (PostgreSQL Async con SQLAlchemy e integraciones de IA) y Capa de Presentación (FastAPI Routers con Inversión de Control)."
    )
    add_image_safely("clean-architecture-diagram.jpg")

    add_heading_2("24.3. Servidor de Despliegue, Hosting Gratuito y Garantía Financiera ($0 USD)")
    doc.add_paragraph(
        "Para garantizar que la Clínica San Juan de Dios NO asuma ningún costo, se utiliza una infraestructura moderna en la nube "
        "con servidores PaaS gratuitos (alternativas modernas a Heroku como Render.com y Oracle Cloud Always Free), base de datos Supabase "
        "y tokens de Inteligencia Artificial gratuitos con Google AI Studio (Gemini 2.0 Flash con hasta 1M de tokens por llamada gratis)."
    )

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    hdr_cells = table.rows[0].cells
    hdr_titles = ["Componente", "Proveedor Gratuito", "Costo", "Garantía Operativa para la Clínica"]
    for i, t in enumerate(hdr_titles):
        hdr_cells[i].text = t
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="2B6CB0"/>')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    infra_data = [
        ("Frontend Web", "Vercel / Cloudflare Pages", "$0 USD (Gratis)", "CDN mundial, HTTPS automático y ancho de banda gratuito."),
        ("Backend API (PaaS)", "Render.com / Oracle Always Free", "$0 USD (Gratis)", "Servidor FastAPI alternativo a Heroku sin costo mensual."),
        ("Base de Datos", "Supabase (PostgreSQL)", "$0 USD (Gratis)", "Postgres gestionado con backups automáticos y 500 MB gratis."),
        ("Inteligencia Artificial", "Google AI Studio (Gemini)", "$0 USD (Gratis)", "1M tokens/llamada gratis para resúmenes terapéuticos y guías."),
        ("Multimedia", "YouTube Embeds + Supabase", "$0 USD (Gratis)", "Streaming de video y audio sin costo de ancho de banda propio.")
    ]

    for comp, prov, cost, guar in infra_data:
        row_cells = table.add_row().cells
        row_cells[0].text = comp
        row_cells[1].text = prov
        row_cells[2].text = cost
        row_cells[2].paragraphs[0].runs[0].bold = True
        row_cells[3].text = guar

    doc.add_paragraph()
    add_image_safely("zero-cost-deployment-diagram.jpg")

    p_guar = doc.add_paragraph()
    p_guar.paragraph_format.left_indent = Inches(0.3)
    p_guar.paragraph_format.right_indent = Inches(0.3)
    r_gt = p_guar.add_run("GARANTÍA FINANCIERA INSTITUCIONAL: ")
    r_gt.bold = True
    r_gt.font.color.rgb = RGBColor(0x9B, 0x2C, 0x2C)
    r_gb = p_guar.add_run(
        "La Clínica San Juan de Dios NO asume ningún costo financiero, de servidores, licenciamiento de software ni mantenimiento técnico. "
        "La infraestructura tecnológica es 100% autosostenible y administrada de forma independiente por el líder técnico del proyecto."
    )
    r_gb.italic = True

    doc.save(output_path)
    print(f"Full proposal document generated successfully at: {output_path}")


if __name__ == "__main__":
    build_full_proposal_docx()
